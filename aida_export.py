"""
AIDA - Eksport moduli
=====================
Tahlil natijalarini grafik + Excel + PDF + Word formatida saqlash.

O'rnatish:
    pip install matplotlib openpyxl reportlab python-docx pandas numpy

Ishlatish:
    from aida_export import AnalysisReport, ReportExporter

    rep = AnalysisReport("DNK tahlili")
    rep.add_table("Nukleotidlar", {"A": 5, "T": 13, "G": 21, "C": 27})
    rep.add_chart(Charts.base_composition({"A": 5, "T": 13, "G": 21, "C": 27}))
    ReportExporter(rep).to_all("natija")   # natija.xlsx, natija.pdf, natija.docx
"""

from __future__ import annotations

import io
import os
from dataclasses import dataclass, field
from datetime import datetime

import matplotlib

matplotlib.use("Agg")  # GUI'siz ishlash — server/fon rejimi uchun
import matplotlib.pyplot as plt
import numpy as np
from matplotlib.backends.backend_agg import FigureCanvasAgg
from matplotlib.figure import Figure


def _subplots(nrows=1, ncols=1, figsize=None, dpi=None, **kw):
    """Thread-xavfsiz subplot — pyplot global holatidan foydalanmaydi.

    pyplot (plt.subplots) global holatga ega va worker oqimida xavfli.
    OO Figure + Agg canvas har oqimda mustaqil ishlaydi.
    """
    fig = Figure(figsize=figsize, dpi=dpi)
    FigureCanvasAgg(fig)
    axes = fig.subplots(nrows, ncols, **kw)
    return fig, axes

# =====================================================================
# USLUB — barcha grafiklar bir xil ko'rinishda
# =====================================================================

PALETTE = {
    "primary": "#2563eb",
    "secondary": "#7c3aed",
    "success": "#059669",
    "warning": "#d97706",
    "danger": "#dc2626",
    "neutral": "#64748b",
    "bases": {"A": "#059669", "T": "#dc2626", "G": "#2563eb", "C": "#d97706", "U": "#dc2626", "N": "#94a3b8"},
}

plt.rcParams.update({
    "figure.dpi": 120,
    "savefig.dpi": 200,
    "savefig.bbox": "tight",
    "font.family": "sans-serif",
    "font.size": 10,
    "axes.spines.top": False,
    "axes.spines.right": False,
    "axes.grid": True,
    "grid.alpha": 0.25,
    "grid.linestyle": "--",
    "axes.axisbelow": True,
})


@dataclass
class Chart:
    """Grafik — xotirada PNG sifatida saqlanadi."""

    title: str
    png: bytes
    caption: str = ""

    def save(self, path: str):
        with open(path, "wb") as f:
            f.write(self.png)
        return path


def _fig_to_chart(fig, title: str, caption: str = "") -> Chart:
    buf = io.BytesIO()
    fig.savefig(buf, format="png")
    return Chart(title=title, png=buf.getvalue(), caption=caption)


# =====================================================================
# GRAFIKLAR KUTUBXONASI
# =====================================================================


class Charts:
    """Har bir tahlil turi uchun tayyor grafik."""

    @staticmethod
    def base_composition(counts: dict, title: str = "Nukleotid tarkibi") -> Chart:
        """Nukleotidlar taqsimoti — ustunli + doiraviy."""
        fig, (ax1, ax2) = _subplots(1, 2, figsize=(10, 4))
        bases = list(counts.keys())
        values = list(counts.values())
        colors = [PALETTE["bases"].get(b, PALETTE["neutral"]) for b in bases]

        ax1.bar(bases, values, color=colors, edgecolor="white", linewidth=1.5)
        ax1.set_ylabel("Soni")
        ax1.set_title("Nukleotidlar soni")
        for i, v in enumerate(values):
            ax1.text(i, v, str(v), ha="center", va="bottom", fontweight="bold")

        ax2.pie(values, labels=bases, colors=colors, autopct="%1.1f%%",
                wedgeprops={"edgecolor": "white", "linewidth": 2})
        ax2.set_title("Foizda")

        fig.suptitle(title, fontsize=13, fontweight="bold")
        return _fig_to_chart(fig, title, f"Jami {sum(values)} nukleotid.")

    @staticmethod
    def gc_content_window(seq: str, window: int = 50, title: str = "GC tarkibi") -> Chart:
        """GC tarkibi sekvensiya bo'ylab — CpG orollari va gen zonalarini ko'rsatadi."""
        s = seq.upper().replace("U", "T")
        if len(s) < window:
            window = max(3, len(s) // 4)

        positions, gc_values = [], []
        for i in range(0, len(s) - window + 1, max(1, window // 5)):
            w = s[i : i + window]
            gc_values.append(100 * (w.count("G") + w.count("C")) / len(w))
            positions.append(i + window // 2)

        fig, ax = _subplots(figsize=(10, 3.5))
        ax.plot(positions, gc_values, color=PALETTE["primary"], linewidth=2)
        ax.fill_between(positions, gc_values, alpha=0.15, color=PALETTE["primary"])
        mean_gc = float(np.mean(gc_values))
        ax.axhline(mean_gc, color=PALETTE["danger"], linestyle="--", linewidth=1.5,
                   label=f"O'rtacha {mean_gc:.1f}%")
        ax.set_xlabel("Pozitsiya (nt)")
        ax.set_ylabel("GC %")
        ax.set_title(title, fontweight="bold")
        ax.legend()
        return _fig_to_chart(fig, title, f"{window} nt oyna bo'yicha. O'rtacha GC: {mean_gc:.1f}%.")

    @staticmethod
    def hardy_weinberg(observed: dict, expected: dict, chi2: float, p_value: float) -> Chart:
        """Hardy-Weinberg: kuzatilgan va kutilgan genotiplar."""
        labels = list(observed.keys())
        x = np.arange(len(labels))
        w = 0.35

        fig, ax = _subplots(figsize=(8, 4.5))
        ax.bar(x - w / 2, [observed[k] for k in labels], w, label="Kuzatilgan",
               color=PALETTE["primary"], edgecolor="white", linewidth=1.5)
        ax.bar(x + w / 2, [expected[k] for k in labels], w, label="Kutilgan (HW)",
               color=PALETTE["neutral"], edgecolor="white", linewidth=1.5, alpha=0.8)

        ax.set_xticks(x)
        ax.set_xticklabels(labels)
        ax.set_ylabel("Individlar soni")
        verdict = "muvozanatda" if p_value > 0.05 else "muvozanatdan chetlashgan"
        ax.set_title(f"Hardy-Weinberg testi — {verdict}", fontweight="bold")
        ax.legend()
        ax.text(0.98, 0.95, f"$\\chi^2$ = {chi2:.2f}\np = {p_value:.4f}",
                transform=ax.transAxes, ha="right", va="top",
                bbox={"facecolor": "white", "edgecolor": PALETTE["neutral"], "alpha": 0.9})
        return _fig_to_chart(fig, "Hardy-Weinberg", f"Xi-kvadrat {chi2:.2f}, p = {p_value:.4f}.")

    @staticmethod
    def allele_frequencies(freqs: dict, title: str = "Allel chastotalari") -> Chart:
        """Allel chastotalari — gorizontal ustunlar."""
        fig, ax = _subplots(figsize=(8, max(2.5, 0.6 * len(freqs))))
        alleles = list(freqs.keys())
        values = list(freqs.values())
        colors = [PALETTE["bases"].get(a, PALETTE["primary"]) for a in alleles]

        bars = ax.barh(alleles, values, color=colors, edgecolor="white", linewidth=1.5)
        ax.set_xlabel("Chastota")
        ax.set_xlim(0, 1)
        ax.set_title(title, fontweight="bold")
        for bar, v in zip(bars, values):
            ax.text(v + 0.02, bar.get_y() + bar.get_height() / 2, f"{v:.3f}",
                    va="center", fontweight="bold")
        return _fig_to_chart(fig, title)

    @staticmethod
    def manhattan_plot(results: list[dict], threshold: float = 5e-8) -> Chart:
        """Manhattan grafigi — GWAS'ning asosiy vizualizatsiyasi.

        results: [{'chromosome': '1', 'position': 12345, 'p_value': 0.001}, ...]
        """
        fig, ax = _subplots(figsize=(11, 4.5))

        by_chrom = {}
        for r in results:
            by_chrom.setdefault(str(r["chromosome"]), []).append(r)

        def chrom_key(c):
            return (0, int(c)) if c.isdigit() else (1, c)

        offset, ticks, tick_labels = 0, [], []
        colors = [PALETTE["primary"], PALETTE["neutral"]]

        for i, chrom in enumerate(sorted(by_chrom, key=chrom_key)):
            variants = sorted(by_chrom[chrom], key=lambda r: r["position"])
            xs = [offset + v["position"] for v in variants]
            ys = [-np.log10(max(v["p_value"], 1e-300)) for v in variants]
            ax.scatter(xs, ys, s=8, color=colors[i % 2], alpha=0.7)
            ticks.append(offset + (max(xs) - min(xs)) / 2 if xs else offset)
            tick_labels.append(chrom)
            offset = max(xs) + 1 if xs else offset

        ax.axhline(-np.log10(threshold), color=PALETTE["danger"], linestyle="--",
                   linewidth=1.5, label=f"Genom chegarasi (p = {threshold:.0e})")
        ax.axhline(-np.log10(1e-5), color=PALETTE["warning"], linestyle=":",
                   linewidth=1.2, label="Taxminiy chegara (p = 1e-5)")

        ax.set_xticks(ticks)
        ax.set_xticklabels(tick_labels, fontsize=8)
        ax.set_xlabel("Xromosoma")
        ax.set_ylabel("$-\\log_{10}(p)$")
        ax.set_title("Manhattan grafigi — GWAS natijalari", fontweight="bold")
        ax.legend(loc="upper right", fontsize=8)
        ax.grid(axis="x", visible=False)

        hits = sum(1 for r in results if r["p_value"] < threshold)
        return _fig_to_chart(fig, "Manhattan grafigi",
                             f"{len(results):,} variant, {hits} tasi chegaradan o'tdi.")

    @staticmethod
    def qq_plot(p_values: list[float]) -> Chart:
        """QQ grafigi — GWAS sifat nazorati (populyatsiya stratifikatsiyasini ko'rsatadi)."""
        p = np.array([max(x, 1e-300) for x in p_values])
        n = len(p)
        observed = -np.log10(np.sort(p))
        expected = -np.log10(np.arange(1, n + 1) / (n + 1))

        # Genomik inflyatsiya koeffitsienti (lambda)
        from scipy import stats as sps
        chi2_obs = sps.chi2.ppf(1 - p, df=1)
        lambda_gc = np.median(chi2_obs) / sps.chi2.ppf(0.5, df=1)

        fig, ax = _subplots(figsize=(5.5, 5.5))
        ax.scatter(expected, observed, s=10, color=PALETTE["primary"], alpha=0.7)
        lim = max(expected.max(), observed.max()) * 1.05
        ax.plot([0, lim], [0, lim], color=PALETTE["danger"], linestyle="--", linewidth=1.5)
        ax.set_xlabel("Kutilgan $-\\log_{10}(p)$")
        ax.set_ylabel("Kuzatilgan $-\\log_{10}(p)$")
        ax.set_title("QQ grafigi", fontweight="bold")
        ax.text(0.05, 0.95, f"$\\lambda_{{GC}}$ = {lambda_gc:.3f}", transform=ax.transAxes,
                va="top", bbox={"facecolor": "white", "edgecolor": PALETTE["neutral"]})

        note = "Lambda 1 ga yaqin — yaxshi." if 0.95 <= lambda_gc <= 1.05 else \
               "Lambda 1 dan uzoq — populyatsiya stratifikatsiyasi bo'lishi mumkin."
        return _fig_to_chart(fig, "QQ grafigi", f"Lambda = {lambda_gc:.3f}. {note}")

    @staticmethod
    def fst_heatmap(matrix: np.ndarray, labels: list[str]) -> Chart:
        """Fst matritsasi — populyatsiyalar orasidagi genetik masofa."""
        fig, ax = _subplots(figsize=(1.2 * len(labels) + 2, 1.0 * len(labels) + 1.5))
        im = ax.imshow(matrix, cmap="YlOrRd", vmin=0)

        ax.set_xticks(range(len(labels)))
        ax.set_yticks(range(len(labels)))
        ax.set_xticklabels(labels, rotation=45, ha="right")
        ax.set_yticklabels(labels)
        ax.grid(visible=False)

        for i in range(len(labels)):
            for j in range(len(labels)):
                val = matrix[i, j]
                color = "white" if val > matrix.max() * 0.6 else "black"
                ax.text(j, i, f"{val:.3f}", ha="center", va="center", color=color, fontsize=9)

        fig.colorbar(im, ax=ax, label="$F_{ST}$", shrink=0.8)
        ax.set_title("Populyatsiyalar orasidagi $F_{ST}$", fontweight="bold", pad=15)
        return _fig_to_chart(fig, "Fst matritsasi",
                             "0 = bir xil, 0.05 dan past = kam farq, 0.15 dan yuqori = katta farq.")

    @staticmethod
    def primer_tm(primers: list[dict]) -> Chart:
        """Primerlar Tm va GC solishtiruvi — PCR dizayni uchun."""
        fig, (ax1, ax2) = _subplots(1, 2, figsize=(11, 4))
        names = [p.get("primer", f"P{i+1}")[:12] for i, p in enumerate(primers)]
        tms = [p["tm"] for p in primers]
        gcs = [p["gc_percent"] for p in primers]

        colors_tm = [PALETTE["success"] if 55 <= t <= 65 else PALETTE["danger"] for t in tms]
        ax1.bar(names, tms, color=colors_tm, edgecolor="white", linewidth=1.5)
        ax1.axhspan(55, 65, alpha=0.12, color=PALETTE["success"], label="Ideal 55-65°C")
        ax1.set_ylabel("Tm (°C)")
        ax1.set_title("Erish harorati")
        ax1.legend(fontsize=8)
        ax1.tick_params(axis="x", rotation=45)

        colors_gc = [PALETTE["success"] if 40 <= g <= 60 else PALETTE["danger"] for g in gcs]
        ax2.bar(names, gcs, color=colors_gc, edgecolor="white", linewidth=1.5)
        ax2.axhspan(40, 60, alpha=0.12, color=PALETTE["success"], label="Ideal 40-60%")
        ax2.set_ylabel("GC (%)")
        ax2.set_title("GC tarkibi")
        ax2.legend(fontsize=8)
        ax2.tick_params(axis="x", rotation=45)

        fig.suptitle("Primer sifati", fontsize=13, fontweight="bold")
        return _fig_to_chart(fig, "Primer sifati", "Yashil = ideal oraliqda, qizil = tuzatish kerak.")

    @staticmethod
    def codon_usage(freqs: dict, top: int = 20) -> Chart:
        """Kodon ishlatilishi — eng ko'p uchraydiganlari."""
        items = sorted(freqs.items(), key=lambda x: x[1], reverse=True)[:top]
        fig, ax = _subplots(figsize=(10, 4))
        ax.bar([c for c, _ in items], [v for _, v in items],
               color=PALETTE["secondary"], edgecolor="white", linewidth=1.2)
        ax.set_ylabel("Chastota")
        ax.set_xlabel("Kodon")
        ax.set_title(f"Eng ko'p uchraydigan {len(items)} kodon", fontweight="bold")
        ax.tick_params(axis="x", rotation=90)
        return _fig_to_chart(fig, "Kodon ishlatilishi")

    @staticmethod
    def prs_distribution(user_score: float, population: list[float] | None = None) -> Chart:
        """PRS taqsimoti — foydalanuvchi populyatsiyaga nisbatan qayerda."""
        if population is None:
            population = list(np.random.normal(0, 1, 10000))

        pop = np.array(population)
        percentile = float((pop < user_score).mean() * 100)

        fig, ax = _subplots(figsize=(9, 4))
        ax.hist(pop, bins=60, color=PALETTE["neutral"], alpha=0.55, edgecolor="white")
        ax.axvline(user_score, color=PALETTE["danger"], linewidth=2.5,
                   label=f"Sizning ballingiz ({percentile:.0f}-persentil)")
        ax.set_xlabel("Poligenik xavf balli (PRS)")
        ax.set_ylabel("Odamlar soni")
        ax.set_title("PRS populyatsiya taqsimotida", fontweight="bold")
        ax.legend()
        return _fig_to_chart(
            fig, "PRS taqsimoti",
            f"{percentile:.0f}-persentil. PRS qiyosiy ko'rsatkich — mutlaq xavf emas "
            f"va yolg'iz klinik qaror uchun ishlatilmaydi.",
        )

    @staticmethod
    def protein_properties(props: dict) -> Chart:
        """Oqsil ikkilamchi strukturasi va xossalari."""
        fig, (ax1, ax2) = _subplots(1, 2, figsize=(10, 4))

        ss = props.get("secondary_structure", {})
        if ss:
            ax1.pie(list(ss.values()), labels=["Spiral (helix)", "Burilish (turn)", "Varaq (sheet)"],
                    colors=[PALETTE["primary"], PALETTE["warning"], PALETTE["success"]],
                    autopct="%1.1f%%", wedgeprops={"edgecolor": "white", "linewidth": 2})
            ax1.set_title("Ikkilamchi struktura")

        metrics = {
            "Og'irlik (kDa)": props.get("molecular_weight_kda", 0),
            "pI": props.get("isoelectric_point", 0),
            "GRAVY": props.get("gravy", 0),
            "Barqarorsizlik": props.get("instability_index", 0),
        }
        colors = [PALETTE["primary"]] * 3 + [
            PALETTE["success"] if metrics["Barqarorsizlik"] < 40 else PALETTE["danger"]
        ]
        ax2.barh(list(metrics.keys()), list(metrics.values()), color=colors,
                 edgecolor="white", linewidth=1.5)
        ax2.set_title("Xossalari")
        for i, v in enumerate(metrics.values()):
            ax2.text(v, i, f" {v:.2f}", va="center", fontweight="bold")

        fig.suptitle("Oqsil tahlili", fontsize=13, fontweight="bold")
        stable = "barqaror" if props.get("stable") else "barqaror emas"
        return _fig_to_chart(fig, "Oqsil tahlili",
                             f"Barqarorsizlik indeksi 40 dan past bo'lsa oqsil barqaror. Bu oqsil {stable}.")


# =====================================================================
# HISOBOT KONTEYNERI
# =====================================================================


@dataclass
class Section:
    heading: str
    text: str = ""
    table: dict | list | None = None
    table_headers: list[str] | None = None
    chart: Chart | None = None


@dataclass
class AnalysisReport:
    """Barcha tahlil natijalarini yig'adigan konteyner."""

    title: str
    subtitle: str = ""
    author: str = "AIDA"
    created: datetime = field(default_factory=datetime.now)
    sections: list[Section] = field(default_factory=list)
    disclaimer: str = (
        "Bu hisobot AIDA tomonidan avtomatik yaratilgan va faqat ma'lumot uchun. "
        "U tibbiy tashxis emas. Genetik natijalar bo'yicha genetik maslahatchi yoki "
        "shifokor bilan maslahatlashing."
    )

    def add_section(self, heading: str, text: str = "") -> AnalysisReport:
        self.sections.append(Section(heading=heading, text=text))
        return self

    def add_table(self, heading: str, data: dict | list, headers: list[str] | None = None,
                  text: str = "") -> AnalysisReport:
        self.sections.append(Section(heading=heading, text=text, table=data, table_headers=headers))
        return self

    def add_chart(self, chart: Chart, heading: str | None = None, text: str = "") -> AnalysisReport:
        self.sections.append(Section(heading=heading or chart.title, text=text, chart=chart))
        return self

    @staticmethod
    def _rows(section: Section) -> tuple[list[str], list[list]]:
        """Jadval ma'lumotini (sarlavhalar, qatorlar) ko'rinishiga keltirish."""
        data = section.table
        if isinstance(data, dict):
            headers = section.table_headers or ["Ko'rsatkich", "Qiymat"]
            rows = [[str(k), v] for k, v in data.items()]
        elif isinstance(data, list) and data and isinstance(data[0], dict):
            headers = section.table_headers or list(data[0].keys())
            rows = [[r.get(h, "") for h in headers] for r in data]
        elif isinstance(data, list):
            headers = section.table_headers or [f"Ustun {i+1}" for i in range(len(data[0]))]
            rows = [list(r) for r in data]
        else:
            headers, rows = [], []
        return headers, rows


# =====================================================================
# EKSPORT
# =====================================================================


class ReportExporter:
    """Hisobotni Excel, PDF va Word formatlariga chiqarish."""

    def __init__(self, report: AnalysisReport):
        self.r = report

    # ---------------- EXCEL ----------------
    def to_xlsx(self, path: str) -> str:
        from openpyxl import Workbook
        from openpyxl.drawing.image import Image as XLImage
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.utils import get_column_letter

        if not path.endswith(".xlsx"):
            path += ".xlsx"

        wb = Workbook()
        ws = wb.active
        ws.title = "Xulosa"

        title_font = Font(name="Arial", size=16, bold=True, color="1E3A8A")
        head_font = Font(name="Arial", size=11, bold=True, color="FFFFFF")
        head_fill = PatternFill("solid", fgColor="2563EB")
        body_font = Font(name="Arial", size=10)
        thin = Side(style="thin", color="CBD5E1")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        # Sarlavha varag'i
        ws["A1"] = self.r.title
        ws["A1"].font = title_font
        ws["A2"] = self.r.subtitle
        ws["A2"].font = Font(name="Arial", size=11, italic=True, color="64748B")
        ws["A3"] = f"Yaratildi: {self.r.created:%d.%m.%Y %H:%M} · {self.r.author}"
        ws["A3"].font = Font(name="Arial", size=9, color="64748B")

        ws["A5"] = "OGOHLANTIRISH"
        ws["A5"].font = Font(name="Arial", size=10, bold=True, color="B91C1C")
        ws["A6"] = self.r.disclaimer
        ws["A6"].font = Font(name="Arial", size=9, color="B91C1C")
        ws["A6"].alignment = Alignment(wrap_text=True, vertical="top")
        ws.merge_cells("A6:F8")
        ws.column_dimensions["A"].width = 32
        for col in "BCDEF":
            ws.column_dimensions[col].width = 16

        # Mundarija
        ws["A10"] = "Mundarija"
        ws["A10"].font = Font(name="Arial", size=12, bold=True)
        row = 11
        for i, s in enumerate(self.r.sections, 1):
            ws.cell(row=row, column=1, value=f"{i}. {s.heading}").font = body_field = body_font
            row += 1

        # Har bir bo'lim uchun alohida varaq
        used_names = {"Xulosa"}
        for i, section in enumerate(self.r.sections, 1):
            name = f"{i}. {section.heading}"[:31]
            while name in used_names:
                name = f"{name[:28]}_{i}"
            used_names.add(name)
            sh = wb.create_sheet(name)

            sh["A1"] = section.heading
            sh["A1"].font = Font(name="Arial", size=14, bold=True, color="1E3A8A")
            r_idx = 3

            if section.text:
                sh.cell(row=r_idx, column=1, value=section.text).font = body_font
                sh.cell(row=r_idx, column=1).alignment = Alignment(wrap_text=True, vertical="top")
                sh.merge_cells(start_row=r_idx, start_column=1, end_row=r_idx + 1, end_column=6)
                r_idx += 3

            if section.table is not None:
                headers, rows = AnalysisReport._rows(section)
                for c, h in enumerate(headers, 1):
                    cell = sh.cell(row=r_idx, column=c, value=h)
                    cell.font, cell.fill, cell.border = head_font, head_fill, border
                    cell.alignment = Alignment(horizontal="center")
                r_idx += 1

                first_data_row = r_idx
                for row_data in rows:
                    for c, v in enumerate(row_data, 1):
                        cell = sh.cell(row=r_idx, column=c,
                                       value=v if isinstance(v, (int, float)) else str(v))
                        cell.font, cell.border = body_font, border
                        if isinstance(v, float):
                            cell.number_format = "0.0000"
                    r_idx += 1

                # Raqamli ustunlar uchun jami — formula bilan, qattiq raqam emas
                if rows and len(rows) > 1:
                    for c in range(1, len(headers) + 1):
                        if all(isinstance(r[c - 1], (int, float)) for r in rows):
                            col = get_column_letter(c)
                            cell = sh.cell(row=r_idx, column=c,
                                           value=f"=SUM({col}{first_data_row}:{col}{r_idx-1})")
                            cell.font = Font(name="Arial", size=10, bold=True)
                            cell.border = border
                    sh.cell(row=r_idx, column=1, value="JAMI").font = Font(
                        name="Arial", size=10, bold=True)
                    r_idx += 1

                # Ustun kengligi
                for c in range(1, len(headers) + 1):
                    width = max([len(str(headers[c - 1]))] +
                                [len(str(r[c - 1])) for r in rows]) + 3
                    sh.column_dimensions[get_column_letter(c)].width = min(width, 45)
                r_idx += 2

            if section.chart:
                img_path = f"/tmp/aida_chart_{i}.png"
                section.chart.save(img_path)
                img = XLImage(img_path)
                img.width, img.height = int(img.width * 0.6), int(img.height * 0.6)
                sh.add_image(img, f"A{r_idx}")
                if section.chart.caption:
                    cap_row = r_idx + 22
                    sh.cell(row=cap_row, column=1, value=section.chart.caption).font = Font(
                        name="Arial", size=9, italic=True, color="64748B")

        wb.save(path)
        return path

    # ---------------- PDF ----------------
    def to_pdf(self, path: str) -> str:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import (Image, PageBreak, Paragraph, SimpleDocTemplate,
                                        Spacer, Table, TableStyle)

        if not path.endswith(".pdf"):
            path += ".pdf"

        doc = SimpleDocTemplate(
            path, pagesize=A4,
            leftMargin=2 * cm, rightMargin=2 * cm, topMargin=2 * cm, bottomMargin=2 * cm,
            title=self.r.title, author=self.r.author,
        )

        ss = getSampleStyleSheet()
        h_title = ParagraphStyle("T", parent=ss["Title"], fontSize=22, spaceAfter=6,
                                 textColor=colors.HexColor("#1E3A8A"))
        h_sub = ParagraphStyle("S", parent=ss["Normal"], fontSize=12, alignment=TA_CENTER,
                               textColor=colors.HexColor("#64748B"), spaceAfter=20)
        h1 = ParagraphStyle("H1", parent=ss["Heading1"], fontSize=15,
                            textColor=colors.HexColor("#1E3A8A"), spaceBefore=14, spaceAfter=8)
        body = ParagraphStyle("B", parent=ss["Normal"], fontSize=10, leading=15, spaceAfter=6)
        caption = ParagraphStyle("C", parent=ss["Normal"], fontSize=8.5, alignment=TA_CENTER,
                                 textColor=colors.HexColor("#64748B"), spaceBefore=4)
        warn = ParagraphStyle("W", parent=ss["Normal"], fontSize=9, leading=13,
                              textColor=colors.HexColor("#B91C1C"),
                              backColor=colors.HexColor("#FEF2F2"),
                              borderPadding=8, borderColor=colors.HexColor("#FCA5A5"), borderWidth=1)

        story = [
            Paragraph(self.r.title, h_title),
            Paragraph(self.r.subtitle or "Genetik tahlil hisoboti", h_sub),
            Paragraph(f"{self.r.author} · {self.r.created:%d.%m.%Y %H:%M}", h_sub),
            Spacer(1, 10),
            Paragraph(f"<b>Ogohlantirish.</b> {self.r.disclaimer}", warn),
            Spacer(1, 18),
        ]

        # Mundarija
        story.append(Paragraph("Mundarija", h1))
        for i, s in enumerate(self.r.sections, 1):
            story.append(Paragraph(f"{i}. {s.heading}", body))
        story.append(PageBreak())

        for i, section in enumerate(self.r.sections, 1):
            story.append(Paragraph(f"{i}. {section.heading}", h1))

            if section.text:
                story.append(Paragraph(section.text, body))
                story.append(Spacer(1, 6))

            if section.table is not None:
                headers, rows = AnalysisReport._rows(section)
                data = [[Paragraph(f"<b>{h}</b>", body) for h in headers]]
                for r in rows:
                    data.append([
                        Paragraph(f"{v:.4f}" if isinstance(v, float) else str(v), body)
                        for v in r
                    ])
                avail = doc.width
                t = Table(data, colWidths=[avail / len(headers)] * len(headers), repeatRows=1)
                t.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2563EB")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1),
                     [colors.white, colors.HexColor("#F8FAFC")]),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CBD5E1")),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]))
                story.append(t)
                story.append(Spacer(1, 10))

            if section.chart:
                img_path = f"/tmp/aida_pdf_chart_{i}.png"
                section.chart.save(img_path)
                from PIL import Image as PILImage
                with PILImage.open(img_path) as im:
                    iw, ih = im.size
                max_w = doc.width
                scale = min(max_w / iw, 1.0)
                story.append(Image(img_path, width=iw * scale, height=ih * scale))
                if section.chart.caption:
                    story.append(Paragraph(section.chart.caption, caption))
                story.append(Spacer(1, 10))

        def footer(canvas, doc_):
            canvas.saveState()
            canvas.setFont("Helvetica", 8)
            canvas.setFillColor(colors.HexColor("#94A3B8"))
            canvas.drawString(2 * cm, 1.2 * cm, f"AIDA · {self.r.created:%d.%m.%Y}")
            canvas.drawRightString(A4[0] - 2 * cm, 1.2 * cm, f"{doc_.page}-bet")
            canvas.setStrokeColor(colors.HexColor("#E2E8F0"))
            canvas.line(2 * cm, 1.6 * cm, A4[0] - 2 * cm, 1.6 * cm)
            canvas.restoreState()

        doc.build(story, onFirstPage=footer, onLaterPages=footer)
        return path

    # ---------------- WORD ----------------
    def to_docx(self, path: str) -> str:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor

        if not path.endswith(".docx"):
            path += ".docx"

        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(10)

        # Sarlavha
        t = doc.add_heading(self.r.title, level=0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sub = doc.add_paragraph(self.r.subtitle or "Genetik tahlil hisoboti")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].italic = True
        sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        meta = doc.add_paragraph(f"{self.r.author} · {self.r.created:%d.%m.%Y %H:%M}")
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.runs[0].font.size = Pt(9)
        meta.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        # Ogohlantirish
        warn = doc.add_paragraph()
        wr = warn.add_run("Ogohlantirish. ")
        wr.bold = True
        wr.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)
        wr2 = warn.add_run(self.r.disclaimer)
        wr2.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)
        wr2.font.size = Pt(9)

        # Mundarija
        doc.add_heading("Mundarija", level=1)
        for i, s in enumerate(self.r.sections, 1):
            doc.add_paragraph(f"{i}. {s.heading}", style="List Number" if False else None)

        doc.add_page_break()

        for i, section in enumerate(self.r.sections, 1):
            doc.add_heading(f"{i}. {section.heading}", level=1)

            if section.text:
                doc.add_paragraph(section.text)

            if section.table is not None:
                headers, rows = AnalysisReport._rows(section)
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = "Light Grid Accent 1"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                hdr = table.rows[0].cells
                for c, h in enumerate(headers):
                    hdr[c].text = str(h)
                    for p in hdr[c].paragraphs:
                        for run in p.runs:
                            run.bold = True

                for row_data in rows:
                    cells = table.add_row().cells
                    for c, v in enumerate(row_data):
                        cells[c].text = f"{v:.4f}" if isinstance(v, float) else str(v)

                doc.add_paragraph()

            if section.chart:
                img_path = f"/tmp/aida_docx_chart_{i}.png"
                section.chart.save(img_path)
                doc.add_picture(img_path, width=Inches(6.2))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

                if section.chart.caption:
                    cap = doc.add_paragraph(section.chart.caption)
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap.runs[0].italic = True
                    cap.runs[0].font.size = Pt(8.5)
                    cap.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        doc.save(path)
        return path

    # ---------------- HAMMASI ----------------
    def to_all(self, base_path: str) -> dict:
        """Uch formatga birdaniga chiqarish."""
        base = os.path.splitext(base_path)[0]
        results = {}
        for fmt, fn in (("xlsx", self.to_xlsx), ("pdf", self.to_pdf), ("docx", self.to_docx)):
            try:
                results[fmt] = fn(f"{base}.{fmt}")
            except Exception as e:
                results[fmt] = f"XATO: {e}"
        return results

    def charts_to_png(self, folder: str) -> list[str]:
        """Faqat grafiklarni PNG sifatida saqlash."""
        os.makedirs(folder, exist_ok=True)
        paths = []
        for i, s in enumerate(self.r.sections, 1):
            if s.chart:
                safe = "".join(c if c.isalnum() else "_" for c in s.chart.title)[:40]
                paths.append(s.chart.save(os.path.join(folder, f"{i:02d}_{safe}.png")))
        return paths


# =====================================================================
# NAMUNA
# =====================================================================

if __name__ == "__main__":
    np.random.seed(42)

    rep = AnalysisReport(
        title="Genetik tahlil hisoboti",
        subtitle="Namuna: DNK sekvensiyasi va populyatsiya tahlili",
    )

    counts = {"A": 5, "T": 13, "G": 21, "C": 27}
    rep.add_table("Nukleotid tarkibi", counts, headers=["Nukleotid", "Soni"],
                  text="Sekvensiyadagi har bir nukleotid soni.")
    rep.add_chart(Charts.base_composition(counts))

    demo_seq = "ATGGCCCTGTGGATGCGCCTCCTGCCCCTGCTGGCGCTGCTGGCCCTCTGGGGACCTGACCCAGCC" * 6
    rep.add_chart(Charts.gc_content_window(demo_seq, window=40))

    rep.add_section("Hardy-Weinberg muvozanati",
                    "Populyatsiyada genotiplar taqsimoti nazariy kutilganga mos keladimi.")
    rep.add_chart(Charts.hardy_weinberg(
        observed={"AA": 320, "Ab": 480, "bb": 200},
        expected={"AA": 313.6, "Ab": 492.8, "bb": 193.6},
        chi2=0.67, p_value=0.4114,
    ))
    rep.add_table("Allel chastotalari", {"A": 0.56, "G": 0.44}, headers=["Allel", "Chastota"])
    rep.add_chart(Charts.allele_frequencies({"A": 0.56, "G": 0.44}))

    gwas = [
        {"chromosome": str(c), "position": p, "p_value": pv}
        for c in range(1, 11)
        for p, pv in zip(
            sorted(np.random.randint(0, 200_000, 220)),
            np.random.uniform(1e-9 if c == 3 else 1e-4, 1, 220),
        )
    ]
    rep.add_chart(Charts.manhattan_plot(gwas))
    rep.add_chart(Charts.qq_plot([g["p_value"] for g in gwas]))

    labels = ["O'zbek", "Qozoq", "Tojik", "Rus"]
    m = np.array([
        [0.000, 0.012, 0.018, 0.045],
        [0.012, 0.000, 0.021, 0.048],
        [0.018, 0.021, 0.000, 0.039],
        [0.045, 0.048, 0.039, 0.000],
    ])
    rep.add_chart(Charts.fst_heatmap(m, labels))

    rep.add_chart(Charts.primer_tm([
        {"primer": "ATGGCCCTGTGGATGCGCC", "tm": 59.7, "gc_percent": 68.4},
        {"primer": "GCTAGCTAGCTTAAGCTAG", "tm": 52.1, "gc_percent": 47.4},
        {"primer": "CGCGCGGCGGCCGCGGCGC", "tm": 72.3, "gc_percent": 89.5},
    ]))

    rep.add_chart(Charts.prs_distribution(user_score=1.4))

    rep.add_chart(Charts.protein_properties({
        "molecular_weight_kda": 24.6, "isoelectric_point": 6.8, "gravy": -0.35,
        "instability_index": 32.4, "stable": True,
        "secondary_structure": {"helix": 0.34, "turn": 0.22, "sheet": 0.28},
    }))

    out = ReportExporter(rep).to_all("/tmp/aida_demo")
    for fmt, p in out.items():
        print(f"{fmt.upper():5} -> {p}")
