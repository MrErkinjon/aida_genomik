"""
AIDA — RPC / ANOVA tahlil moduli
================================
80 RIL uchun har bir trait (RPC ustuni) bo'yicha bir tomonlama ANOVA,
LSD, Mean, SD/SE, ahamiyat harflari (compact letter display), tolerantlik
tasnifi, publikatsiya darajasidagi bar grafiklar va Excel/Word eksport.

Statistik asos:
    - ANOVA: trait ~ Genotype (one-way), replikatlar ichida.
    - LSD = t(α/2, df_error) · √(2·MSE/r)
    - SE = √(MSE/r)
    - CLD: LSD asosida juftlik taqqoslashdan harflar (a, b, c ...)

MUHIM: RPC = ((stress − control)/control)·100. Ijobiy/kamroq manfiy RPC
o'sish traitlari uchun ko'proq tolerantlikni bildiradi (yo'nalish trait'ga bog'liq).
"""

from __future__ import annotations

from dataclasses import dataclass, field

import matplotlib

matplotlib.use("Agg")
import numpy as np
import pandas as pd
from matplotlib.backends.backend_agg import FigureCanvasAgg
from matplotlib.figure import Figure
from scipy import stats


def _subplots(nrows=1, ncols=1, figsize=None, dpi=None, **kw):
    """Thread-xavfsiz subplot (pyplot global holatisiz — worker uchun)."""
    fig = Figure(figsize=figsize, dpi=dpi)
    FigureCanvasAgg(fig)
    return fig, fig.subplots(nrows, ncols, **kw)

# =====================================================================
# 0. YUKLASH VA USTUNLARNI ANIQLASH
# =====================================================================

GENO_ALIASES = {"genotype", "geno", "line", "ril", "entry", "name"}
REP_ALIASES = {"rep", "replication", "replicate", "block", "r"}


def load_rpc(path: str, sheet=0) -> pd.DataFrame:
    """RPC Excel/CSV faylini o'qiydi."""
    if str(path).lower().endswith(".csv"):
        return pd.read_csv(path)
    return pd.read_excel(path, sheet_name=sheet)


def detect_columns(df: pd.DataFrame) -> tuple[str, str | None, list[str]]:
    """(genotype_ustuni, rep_ustuni|None, trait_ustunlari) ni aniqlaydi."""
    geno_col, rep_col = None, None
    for c in df.columns:
        lc = str(c).strip().lower()
        if geno_col is None and lc in GENO_ALIASES:
            geno_col = c
        elif rep_col is None and lc in REP_ALIASES:
            rep_col = c
    if geno_col is None:
        geno_col = df.columns[0]  # birinchi ustun — genotip deb faraz
    traits = [c for c in df.columns
              if c not in (geno_col, rep_col)
              and pd.api.types.is_numeric_dtype(df[c])]
    return geno_col, rep_col, traits


# =====================================================================
# 1. BIR TOMONLAMA ANOVA + LSD
# =====================================================================


@dataclass
class AnovaResult:
    trait: str
    n_genotypes: int
    n_obs: int
    reps: float             # o'rtacha replikat soni
    grand_mean: float
    f_value: float
    p_value: float
    df_between: int
    df_within: int
    ms_between: float
    ms_error: float         # MSE
    lsd_05: float
    lsd_01: float
    cv_percent: float
    means: pd.DataFrame     # genotip | mean | sd | se | n | letters | group
    sig_code: str

    def summary_row(self) -> dict:
        return {
            "Trait": self.trait,
            "Genotiplar": self.n_genotypes,
            "N": self.n_obs,
            "Grand mean": round(self.grand_mean, 3),
            "F": round(self.f_value, 3),
            "P": self.p_value,
            "Sig": self.sig_code,
            "LSD(0.05)": round(self.lsd_05, 3),
            "LSD(0.01)": round(self.lsd_01, 3),
            "MSE": round(self.ms_error, 3),
            "CV%": round(self.cv_percent, 2),
        }


def sig_code(p: float) -> str:
    return "***" if p < 0.001 else "**" if p < 0.01 else "*" if p < 0.05 else "ns"


def one_way_anova(df: pd.DataFrame, geno_col: str, trait: str) -> AnovaResult:
    """Bitta trait uchun to'liq bir tomonlama ANOVA + LSD + CLD."""
    sub = df[[geno_col, trait]].dropna()
    grouped = sub.groupby(geno_col)[trait]
    group_means = grouped.mean()
    group_n = grouped.count()
    group_sd = grouped.std(ddof=1)

    k = len(group_means)               # genotiplar soni
    N = int(group_n.sum())             # umumiy kuzatuv
    grand = sub[trait].mean()

    ss_between = float((group_n * (group_means - grand) ** 2).sum())
    ss_total = float(((sub[trait] - grand) ** 2).sum())
    ss_within = ss_total - ss_between

    df_b = k - 1
    df_w = N - k
    ms_b = ss_between / df_b if df_b > 0 else float("nan")
    mse = ss_within / df_w if df_w > 0 else float("nan")
    F = ms_b / mse if mse > 0 else float("nan")
    p = float(stats.f.sf(F, df_b, df_w)) if np.isfinite(F) else float("nan")

    r = float(group_n.mean())          # o'rtacha replikat (muvozanatsizlik uchun)
    t05 = stats.t.ppf(0.975, df_w) if df_w > 0 else float("nan")
    t01 = stats.t.ppf(0.995, df_w) if df_w > 0 else float("nan")
    lsd05 = t05 * np.sqrt(2 * mse / r)
    lsd01 = t01 * np.sqrt(2 * mse / r)
    se = np.sqrt(mse / r)
    cv = 100 * np.sqrt(mse) / abs(grand) if grand != 0 else float("nan")

    means = pd.DataFrame({
        "genotype": group_means.index.astype(str),
        "mean": group_means.values,
        "sd": group_sd.values,
        "se": se,
        "n": group_n.values,
    }).sort_values("mean", ascending=False).reset_index(drop=True)

    letters = compact_letters(means["mean"].values, lsd05)
    means["letters"] = letters
    means["group"] = [classify_tolerance(m, grand, sub[trait].std(ddof=1))
                      for m in means["mean"].values]

    return AnovaResult(
        trait=trait, n_genotypes=k, n_obs=N, reps=r, grand_mean=grand,
        f_value=F, p_value=p, df_between=df_b, df_within=df_w,
        ms_between=ms_b, ms_error=mse, lsd_05=lsd05, lsd_01=lsd01,
        cv_percent=cv, means=means, sig_code=sig_code(p),
    )


def compact_letters(means_desc: np.ndarray, lsd: float) -> list[str]:
    """Compact Letter Display — kamayish tartibida saralangan o'rtachalar uchun.

    Bir xil LSD (muvozanatli dizayn) da 'ahamiyatsiz farq' contiguous interval
    hosil qiladi; maksimal bloklarga harf beriladi.
    """
    n = len(means_desc)
    if n == 0 or not np.isfinite(lsd):
        return [""] * n
    # reach[i] — means[i] dan LSD ichidagi eng uzoq indeks
    reach = [0] * n
    j = 0
    for i in range(n):
        if j < i:
            j = i
        while j + 1 < n and (means_desc[i] - means_desc[j + 1]) <= lsd:
            j += 1
        reach[i] = j

    groups: list[set[int]] = [set() for _ in range(n)]
    letter_ord = 0
    prev_reach = -1
    for i in range(n):
        if reach[i] > prev_reach:
            for k in range(i, reach[i] + 1):
                groups[k].add(letter_ord)
            letter_ord += 1
            prev_reach = reach[i]

    def to_letters(idxs: set[int]) -> str:
        return "".join(chr(ord("a") + o) for o in sorted(idxs))

    return [to_letters(g) for g in groups]


def classify_tolerance(mean_val: float, grand: float, sd: float) -> str:
    """RPC o'rtachasiga qarab tolerantlik guruhi (yo'nalish: yuqori RPC = tolerant)."""
    if not np.isfinite(sd) or sd == 0:
        return "o'rtacha"
    if mean_val >= grand + 0.5 * sd:
        return "chidamli"
    if mean_val <= grand - 0.5 * sd:
        return "chidamsiz"
    return "o'rtacha"


def analyze_all(df: pd.DataFrame, geno_col: str, traits: list[str]) -> list[AnovaResult]:
    return [one_way_anova(df, geno_col, t) for t in traits]


# =====================================================================
# 2. GRAFIK (publikatsiya uslubi)
# =====================================================================

_C = {"bar": "#2563eb", "err": "#1e293b", "res": "#059669",
      "mod": "#d97706", "sus": "#dc2626", "line": "#64748b"}


def bar_chart(res: AnovaResult, show_letters: bool = True, top: int | None = None) -> bytes:
    """Genotip o'rtacha RPC bar grafigi — xatolik chizig'i (SE) + CLD harflari.

    top: faqat eng yuqori+eng past N tasini ko'rsatish (80 ta ko'p bo'lsa).
    """
    m = res.means
    if top and len(m) > 2 * top:
        m = pd.concat([m.head(top), m.tail(top)])
        subtitle = f"(eng yuqori {top} va eng past {top} genotip)"
    else:
        subtitle = f"(barcha {len(m)} genotip)"

    colors = {"chidamli": _C["res"], "o'rtacha": _C["mod"], "chidamsiz": _C["sus"]}
    bar_colors = [colors[g] for g in m["group"]]

    fig, ax = _subplots(figsize=(max(8, len(m) * 0.16), 4.6), dpi=120)
    x = np.arange(len(m))
    ax.bar(x, m["mean"], yerr=m["se"], color=bar_colors, edgecolor="white",
           linewidth=0.4, error_kw={"ecolor": _C["err"], "elinewidth": 0.8, "capsize": 2})
    ax.axhline(res.grand_mean, color=_C["line"], linestyle="--", linewidth=1,
               label=f"Grand mean = {res.grand_mean:.2f}")

    if show_letters and len(m) <= 60:
        for xi, (_, row) in zip(x, m.iterrows()):
            y = row["mean"] + np.sign(row["mean"] or 1) * (row["se"] + abs(res.grand_mean) * 0.03)
            ax.text(xi, y, row["letters"], ha="center",
                    va="bottom" if row["mean"] >= 0 else "top", fontsize=6, color="#334155")

    ax.set_xticks(x)
    ax.set_xticklabels(m["genotype"], rotation=90, fontsize=6)
    ax.set_ylabel(f"{res.trait} (RPC %)")
    ax.set_title(f"{res.trait} — genotiplar bo'yicha RPC  {subtitle}\n"
                 f"F = {res.f_value:.2f}{res.sig_code}, LSD(0.05) = {res.lsd_05:.2f}, CV = {res.cv_percent:.1f}%",
                 fontsize=10, fontweight="bold")
    ax.spines[["top", "right"]].set_visible(False)
    ax.grid(axis="y", alpha=0.2, linestyle="--")
    ax.legend(fontsize=8, loc="best")

    # afsona (guruh ranglari)
    from matplotlib.patches import Patch
    handles = [Patch(facecolor=colors[g], label=lab) for g, lab in
               [("chidamli", "Chidamli"), ("o'rtacha", "O'rtacha"), ("chidamsiz", "Chidamsiz")]]
    ax.legend(handles=[Patch(facecolor=_C["line"], label=f"Grand mean {res.grand_mean:.2f}")] + handles,
              fontsize=7, loc="upper right", ncol=2)

    fig.tight_layout()
    import io
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight")
    return buf.getvalue()


# =====================================================================
# 3. EKSPORT — Excel (har trait alohida varaq) + Word (Results & Discussion)
# =====================================================================


def to_excel(results: list[AnovaResult], path: str) -> str:
    """ANOVA natijalarini Excel'ga: xulosa + har trait uchun varaq."""
    if not path.endswith(".xlsx"):
        path += ".xlsx"
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    summary = pd.DataFrame([r.summary_row() for r in results])
    with pd.ExcelWriter(path, engine="openpyxl") as xw:
        summary.to_excel(xw, sheet_name="Xulosa", index=False)
        for r in results:
            name = _safe_sheet(r.trait)
            tbl = r.means.rename(columns={
                "genotype": "Genotip", "mean": "Mean", "sd": "SD",
                "se": "SE", "n": "N", "letters": "Guruh (LSD)", "group": "Tasnif"})
            tbl.to_excel(xw, sheet_name=name, index=False, startrow=6)
            ws = xw.sheets[name]
            ws["A1"] = f"Trait: {r.trait}"
            ws["A2"] = f"F = {r.f_value:.3f} ({r.sig_code}),  P = {r.p_value:.3e}"
            ws["A3"] = f"LSD(0.05) = {r.lsd_05:.3f},  LSD(0.01) = {r.lsd_01:.3f}"
            ws["A4"] = f"MSE = {r.ms_error:.3f},  CV = {r.cv_percent:.2f}%,  Grand mean = {r.grand_mean:.3f}"
            for row in ("A1", "A2", "A3", "A4"):
                ws[row].font = Font(bold=(row == "A1"), size=12 if row == "A1" else 10)

    # xulosa varag'ini bezash
    from openpyxl import load_workbook
    wb = load_workbook(path)
    ws = wb["Xulosa"]
    head_fill = PatternFill("solid", fgColor="2563EB")
    for c in range(1, summary.shape[1] + 1):
        cell = ws.cell(row=1, column=c)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = head_fill
        cell.alignment = Alignment(horizontal="center")
        ws.column_dimensions[get_column_letter(c)].width = 14
    wb.save(path)
    return path


def to_docx(results: list[AnovaResult], path: str,
            trait_labels: dict[str, str] | None = None) -> str:
    """'Results and Discussion' bo'limi + 19 trait tasnifi (docx)."""
    if not path.endswith(".docx"):
        path += ".docx"
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    trait_labels = trait_labels or {}
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)

    h = doc.add_heading("Results and Discussion", level=1)
    doc.add_paragraph(
        "One-way analysis of variance (ANOVA) was performed for each of the "
        f"{len(results)} relative percentage change (RPC) traits to test for "
        "significant genetic variation among the 80 recombinant inbred lines "
        "(RILs) evaluated under PEG-6000 (15%) induced osmotic stress. Least "
        "significant difference (LSD) at P = 0.05 was used for mean separation, "
        "and lines were grouped into stress-tolerant, moderately tolerant and "
        "stress-susceptible classes based on their RPC values.")

    # umumiy jadval
    sig = sum(1 for r in results if r.p_value < 0.05)
    doc.add_paragraph(
        f"Of the {len(results)} traits, {sig} showed statistically significant "
        f"differences among genotypes (P < 0.05), indicating substantial genetic "
        f"variability in stress response and good potential for selection.")

    for r in results:
        label = trait_labels.get(r.trait, r.trait)
        doc.add_heading(f"{label} ({r.trait})", level=2)
        counts = r.means["group"].value_counts().to_dict()
        nres = counts.get("chidamli", 0)
        nmod = counts.get("o'rtacha", 0)
        nsus = counts.get("chidamsiz", 0)
        top = r.means.head(3)
        bot = r.means.tail(3)
        p = doc.add_paragraph()
        p.add_run(
            f"The trait {label} showed an F-value of {r.f_value:.2f} ({r.sig_code}) "
            f"with LSD₀.₀₅ = {r.lsd_05:.2f} and CV = {r.cv_percent:.1f}%. "
            f"The grand mean RPC was {r.grand_mean:.2f}%. ")
        if r.p_value < 0.05:
            p.add_run(
                "Significant genotypic differences were detected, confirming that "
                "the RILs differ in their physiological response to osmotic stress. ")
        else:
            p.add_run("Differences among genotypes were not statistically significant. ")
        p.add_run(
            f"Based on RPC values, {nres} lines were classified as stress-tolerant, "
            f"{nmod} as moderately tolerant and {nsus} as stress-susceptible. "
            f"The most tolerant lines were {', '.join(top['genotype'])}, whereas "
            f"{', '.join(bot['genotype'])} were the most affected.")

        # tasnif jadvali (qisqa)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Chidamli", "O'rtacha", "Chidamsiz"
        row = table.add_row().cells
        for i, grp in enumerate(["chidamli", "o'rtacha", "chidamsiz"]):
            names = r.means[r.means["group"] == grp]["genotype"].tolist()
            row[i].text = ", ".join(names) if names else "—"

    # cheklovlar
    doc.add_heading("Limitations", level=2)
    doc.add_paragraph(
        "The classification is based solely on RPC magnitude under controlled PEG "
        "conditions; validation under field drought and across multiple environments "
        "is recommended. Trait direction (whether higher RPC implies tolerance) should "
        "be interpreted per trait biology.")

    doc.save(path)
    return path


def _safe_sheet(name: str) -> str:
    s = "".join(c if c.isalnum() or c in " _-" else "_" for c in str(name))
    return s[:31] or "Sheet"


# =====================================================================
# 4. NAMUNA MA'LUMOT (foydalanuvchi fayli bo'lmaganda sinash/demo uchun)
# =====================================================================

RPC_TRAITS = [
    "TPL_RPC", "ShtL_RPC", "RtL_RPC", "TFW_RPC", "FShtW_RPC", "FRtW_RPC",
    "TDW_RPC", "DShtW_RPC", "DRtW_RPC", "FRtW/FShtW_RPC", "DRtW/DShtW_RPC",
    "TWC_RPC", "ShtWC_RPC", "RtWC_RPC", "TWCP_RPC", "ShtWCP_RPC", "RtWCP_RPC",
    "ShtWD_RPC", "RtWD_RPC",
]


def sample_rpc_data(n_lines: int = 80, reps: int = 3, seed: int = 42) -> pd.DataFrame:
    """L_RPC tuzilmasiga o'xshash sun'iy ma'lumot (namoyish uchun)."""
    rng = np.random.default_rng(seed)
    rows = []
    line_effects = {t: rng.normal(0, 12, n_lines) for t in RPC_TRAITS}
    trait_base = {t: rng.uniform(-45, 5) for t in RPC_TRAITS}
    for i in range(n_lines):
        gid = f"RIL{i+1:03d}"
        for rep in range(1, reps + 1):
            row = {"Genotype": gid, "Rep": rep}
            for t in RPC_TRAITS:
                row[t] = trait_base[t] + line_effects[t][i] + rng.normal(0, 6)
            rows.append(row)
    return pd.DataFrame(rows)


if __name__ == "__main__":
    df = sample_rpc_data()
    gcol, rcol, traits = detect_columns(df)
    print(f"Genotip: {gcol}, Rep: {rcol}, traitlar: {len(traits)}")
    results = analyze_all(df, gcol, traits)
    print(f"\n{'Trait':16} {'F':>8} {'P':>10} {'sig':>4} {'LSD05':>8} {'CV%':>6}")
    for r in results[:6]:
        print(f"{r.trait:16} {r.f_value:8.2f} {r.p_value:10.2e} {r.sig_code:>4} "
              f"{r.lsd_05:8.2f} {r.cv_percent:6.1f}")
    print("\nBirinchi trait tasnifi:", results[0].means["group"].value_counts().to_dict())
